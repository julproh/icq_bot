
import numpy as np
import pandas as pd
import xlsxwriter
import docx
def doctable(data, tabletitle, pathfile):
    from docx import Document
    import pandas as pd
    document = Document()
    data = pd.DataFrame(data)  # My input data is in the 2D list form
    document.add_heading(tabletitle)
    table = document.add_table(rows=(data.shape[0]), cols=data.shape[1])  # First row are table headers!
    for i, column in enumerate(data) :
        for row in range(data.shape[0]) :
            table.cell(row, i).text = str(data[column][row])
    document.save(pathfile)
print('Введи количество необходимых варинатов:')
number = int(input())

Data = pd.read_excel('jul.xlsx')

result = []
col = list(Data.columns)
for j in range(number):
    k = np.random.randint(0, 10, 10)
    tasks = []
    for i in range(5):
        tasks.append(col[i] + ' ' + Data.iloc[k[i], i])
    result.append(tasks)


new_names = []
for i in range(1, 6):
    new_names.append('Задача №{}'.format(i))  # переименовка
new_names

New = pd.DataFrame(result)
New.columns = new_names
New['Номер варианта контрольной'] = range(1, number+1)  #добавление стоблца

'''
#запись в эксель файл
writer = pd.ExcelWriter('Jul_outcomes.xlsx', engine='xlsxwriter')
New.to_excel(writer, sheet_name='Sheet1')
writer.save()
'''

for i in New.T:
    doctable(New.T[i],"Вариант "+str(i+1),"Вариант "+str(i+1)+".docx")
